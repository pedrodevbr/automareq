"""
templates.py — Word document template filling and PDF conversion.

Provides:
  - substituir_texto:       Replace placeholders in DOCX body and tables
  - converter_docx_para_pdf: Convert DOCX to PDF via Word automation
  - solicitar_aprovacao_cpv: Fill CPV template, convert to PDF, email for signing
"""

from __future__ import annotations

import datetime
import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)


def substituir_texto(doc, dados: dict) -> None:
    """
    Replace placeholder keys with values in a python-docx Document
    (paragraphs and table cells).
    """
    def _replace_in_paragraph(paragraph, map_dados):
        for chave, valor in map_dados.items():
            if chave in paragraph.text:
                paragraph.text = paragraph.text.replace(chave, str(valor))

    for p in doc.paragraphs:
        _replace_in_paragraph(p, dados)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, dados)


def converter_docx_para_pdf(caminho_docx: str | Path) -> str:
    """
    Convert DOCX to PDF using Word automation (win32com).
    Returns the path of the generated PDF.
    """
    import win32com.client  # lazy import — Windows only

    path_docx = str(Path(caminho_docx).resolve())
    path_pdf = str(Path(caminho_docx).with_suffix(".pdf").resolve())

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    doc = None
    try:
        doc = word.Documents.Open(path_docx)
        doc.SaveAs(path_pdf, FileFormat=17)  # 17 = PDF
    except Exception as exc:
        logger.error("PDF conversion error: %s", exc)
        raise
    finally:
        if doc:
            doc.Close(SaveChanges=False)
        word.Quit()

    return path_pdf


def solicitar_aprovacao_cpv(req: str, resp: str) -> None:
    """
    Fill the CPV declaration template, convert to PDF, and email for signing.
    """
    from docx import Document

    from config.paths import AD_TEMPLATE_DIR
    from config.personnel import CPV_CC_EMAIL, CPV_RECIPIENTS

    dados_preenchimento = {
        "[REQ_NUMBER]": req,
        "[date]": datetime.date.today().strftime("%d/%m/%Y"),
        "[RESPONSAVEL]": resp,
    }

    template_path = AD_TEMPLATE_DIR / "Declaracao_CPV_template.docx"
    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)
    substituir_texto(doc, dados_preenchimento)

    temp_docx = Path(f"Declaracao_{req}.docx")
    doc.save(temp_docx)
    logger.info("DOCX generated: %s", temp_docx)

    arquivo_pdf = converter_docx_para_pdf(temp_docx)

    # Send email
    from core.emitters.stages.send_drafts import enviar_email

    assunto = f"Declaração REQ/SOLPE {req}"
    corpo = (
        f"Prezados,\n\n"
        f"Segue a declaração referente à REQ/SOLPE número {req} para aquisição direta por CPV.\n"
        f"Favor assinar para prosseguir com a aquisição.\n\n"
        f"Atenciosamente,"
    )
    destinatario = CPV_RECIPIENTS.get("teste", list(CPV_RECIPIENTS.values())[0])
    enviar_email(destinatario, assunto, corpo, anexo_path=arquivo_pdf, cc=CPV_CC_EMAIL)

    # Cleanup temp files
    os.remove(temp_docx)
    os.remove(arquivo_pdf)
