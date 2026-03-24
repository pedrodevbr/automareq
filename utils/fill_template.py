import os
import datetime
from pathlib import Path
import win32com.client
from docx import Document

# =============================================================================
# CONFIGURAÇÕES
# =============================================================================
CONFIG = {
    "template": "templates\\AD\\Declaracao_CPV_template.docx",
    "cc_email": "pedrohvb@itaipu.gov.br",
    "opcoes_destinatario": {
        "1": "srvegaa@itaipu.gov.py",
        "2": "borchard@itaipu.gov.br",
        "teste":"pedrohvb@itaipu.gov.br"
    }
}

# =============================================================================
# FUNÇÕES AUXILIARES
# =============================================================================

def substituir_texto(doc, dados):
    """
    Substitui chaves pelos valores no documento (parágrafos e tabelas).
    """
    def _replace_in_paragraph(paragraph, map_dados):
        for chave, valor in map_dados.items():
            if chave in paragraph.text:
                # O uso de 'run' preserva a formatação (negrito, fonte, etc) melhor
                # mas para substituição simples, direto no text resolve a maioria dos casos
                paragraph.text = paragraph.text.replace(chave, str(valor))

    # 1. Substituir no corpo do texto
    for p in doc.paragraphs:
        _replace_in_paragraph(p, dados)

    # 2. Substituir dentro de tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, dados)

def converter_docx_para_pdf(caminho_docx):
    """
    Converte DOCX para PDF usando automação do Word.
    Retorna o caminho do PDF gerado.
    """
    path_docx = str(Path(caminho_docx).resolve())
    path_pdf = str(Path(caminho_docx).with_suffix('.pdf').resolve())
    
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    
    doc = None
    try:
        doc = word.Documents.Open(path_docx)
        # FileFormat 17 = PDF
        doc.SaveAs(path_pdf, FileFormat=17)
    except Exception as e:
        print(f"Erro na conversão PDF: {e}")
        raise
    finally:
        if doc:
            doc.Close(SaveChanges=False)
        # Opcional: word.Quit() fecha o Word inteiro. 
        # Se você costuma ter o Word aberto, comente a linha abaixo.
        word.Quit()
        
    return path_pdf

def enviar_email(destinatario, assunto, corpo, anexo_path=None):
    """
    Envia e-mail via Outlook Desktop.
    """
    try:
        outlook = win32com.client.Dispatch('Outlook.Application')
        mail = outlook.CreateItem(0)
        mail.To = destinatario
        mail.CC = CONFIG["cc_email"]
        mail.Subject = assunto
        mail.Body = corpo

        if anexo_path and Path(anexo_path).exists():
            mail.Attachments.Add(str(Path(anexo_path).resolve()))
        
        mail.Send()
        print(f"✅ E-mail enviado com sucesso para: {destinatario}")
    except Exception as e:
        print(f"❌ Erro ao enviar e-mail: {e}")

def solitar_aprovacao_CPV(req,resp):
    """
    Função placeholder para preencher o CPV na SAP.
    """

    # Dicionário de Substituição (Mapeia o placeholder do Word -> Valor)
    dados_preenchimento = {
        "[REQ_NUMBER]": req,
        "[date]": datetime.date.today().strftime('%d/%m/%Y'),
        "[RESPONSAVEL]":resp
    }
    # 3. Processamento do Arquivo
    try:
        template_path = Path(CONFIG["template"])
        if not template_path.exists():
            raise FileNotFoundError(f"Template não encontrado: {template_path}")

        # Carregar e preencher
        doc = Document(template_path)
        substituir_texto(doc, dados_preenchimento)

        # Salvar temporário
        temp_docx = Path(f"Declaracao_{req}.docx")
        doc.save(temp_docx)
        print(f"Arquivo DOCX gerado: {temp_docx}")

        # Converter para PDF
        print("Convertendo para PDF...")
        arquivo_pdf = converter_docx_para_pdf(temp_docx)

        # 4. Envio de E-mail
        assunto = f"Declaração REQ/SOLPE {req}"
        corpo = (
            f"Prezados,\n\n"
            f"Segue a declaração referente à REQ/SOLPE número {req} para aquisição direta por CPV.\n"
            f"Favor assinar para prosseguir com a aquisição.\n\n"
            f"Atenciosamente,"
        )
        destinatario = CONFIG["opcoes_destinatario"]["teste"]  # Envia para o primeiro destinatário por padrão
        
        enviar_email(destinatario, assunto, corpo, arquivo_pdf)

        # Limpeza (Opcional: remove o docx temporário)
        os.remove(temp_docx)
        os.remove(arquivo_pdf)

    except Exception as e:
        print(f"FATAL ERROR: {e}")
    
# =============================================================================
# FLUXO PRINCIPAL
# =============================================================================

def main():
    print("--- GERADOR DE DECLARAÇÃO REQ/SOLPE ---")
    
    # 1. Coleta de Dados
    req = input("Número da REQ/SOLPE: ")
    # Adicionei os campos extras baseados no seu texto de exemplo
    fornecedor = input("Nome do Fornecedor: ") 
    valor = input("Valor (ex: R$ 1.000,00): ")
    item_desc = input("Descrição do Objeto/Item: ")

    # Dicionário de Substituição (Mapeia o placeholder do Word -> Valor)
    dados_preenchimento = {
        "[REQ_NUMBER]": req,
        "[date]": datetime.date.today().strftime('%d/%m/%Y'),
        "[SUPPLIER]": fornecedor,
        "[REQ_VALUE]": valor,
        "[REQ_ITEM]": item_desc
    }

    # 2. Seleção de Destinatário
    print("\nEscolha o destinatário:")
    print("1 - srvegaa@itaipu.gov.py")
    print("2 - borchard@itaipu.gov.br")
    print("3 - Outro")
    
    escolha = input("Opção: ")
    destinatario = CONFIG["opcoes_destinatario"].get(escolha)
    
    if escolha == "3" or not destinatario:
        destinatario = input("Digite o e-mail do destinatário: ")

    # 3. Processamento do Arquivo
    try:
        template_path = Path(CONFIG["template"])
        if not template_path.exists():
            raise FileNotFoundError(f"Template não encontrado: {template_path}")

        # Carregar e preencher
        doc = Document(template_path)
        substituir_texto(doc, dados_preenchimento)

        # Salvar temporário
        temp_docx = Path(f"Declaracao_{req}.docx")
        doc.save(temp_docx)
        print(f"Arquivo DOCX gerado: {temp_docx}")

        # Converter para PDF
        print("Convertendo para PDF...")
        arquivo_pdf = converter_docx_para_pdf(temp_docx)

        # 4. Envio de E-mail
        assunto = f"Declaração REQ/SOLPE {req}"
        corpo = (
            f"Prezados,\n\n"
            f"Segue a declaração referente à REQ/SOLPE número {req} para aquisição direta por CPV.\n"
            f"Favor assinar para prosseguir com a aquisição.\n\n"
            f"Atenciosamente,"
        )
        
        enviar_email(destinatario, assunto, corpo, arquivo_pdf)

        # Limpeza (Opcional: remove o docx temporário)
        # os.remove(temp_docx)

    except Exception as e:
        print(f"FATAL ERROR: {e}")

if __name__ == "__main__":
    main()